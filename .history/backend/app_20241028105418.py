import json
import os
import shutil
import time
import uuid
from datetime import datetime  # 添加這行
from pathlib import Path
from typing import Dict, List, Optional

import pdfplumber
from config import Config
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.websockets import WebSocket
from pydantic import BaseModel
from rag_utils import (
    add_translated_content_to_vector_store,
    delete_from_vector_store,
    initialize_rag,
    initialize_vector_store,
    query_knowledge_base,
    reset_vector_store,
)
from translation_utils import (
    one_chunk_translate_text,
    translate_and_store_to_knowledge_base,
)


# Pydantic models
class KnowledgeBase(BaseModel):
    name: str
    description: str = ""


class TranslateRequest(BaseModel):
    text: str


class EmbedRequest(BaseModel):
    content: str
    filename: str
    knowledge_base_id: Optional[str] = None


class QueryRequest(BaseModel):
    query: str
    knowledge_base_id: Optional[str] = None
    model_settings: Optional[Dict] = None


class ProgressMessage(BaseModel):
    progress: int


# 新增檔案相關的 Pydantic 模型
class FileInfo(BaseModel):
    id: str
    name: str
    size: int
    type: str
    created_at: str
    category: Optional[str] = None
    tags: List[str] = []


def safely_delete_directory(path: Path):
    """安全地刪除目錄，包括等待和重試機制"""
    max_attempts = 3
    for attempt in range(max_attempts):
        try:
            if path.exists():
                # 先嘗試刪除 sqlite 文件
                sqlite_file = path / "chroma.sqlite3"
                if sqlite_file.exists():
                    sqlite_file.unlink()

                # 等待一下
                time.sleep(1)

                # 刪除整個目錄
                shutil.rmtree(path)
            return True
        except Exception as e:
            print(f"刪除嘗試 {attempt + 1} 失敗: {str(e)}")
            if attempt < max_attempts - 1:
                time.sleep(2)  # 在重試之前等待
            else:
                raise


app = FastAPI()

# CORS 設置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 初始化默認知識庫
vector_store, ffm = initialize_rag()

# WebSocket 連接
active_connections: List[WebSocket] = []


@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    active_connections.append(websocket)
    try:
        while True:
            await websocket.receive_text()
    except:
        active_connections.remove(websocket)


async def broadcast_progress(progress: int):
    for connection in active_connections:
        await connection.send_json({"progress": progress})


def load_knowledge_bases():
    kb_file = Path(Config.CHROMA_PATH) / "knowledge_bases.json"
    if kb_file.exists():
        with open(kb_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return {
        "default": {
            "name": "預設知識庫",
            "description": "預設的知識庫",
            "path": str(Path(Config.CHROMA_PATH) / "default"),
        }
    }


def save_knowledge_bases(knowledge_bases: dict):
    kb_file = Path(Config.CHROMA_PATH) / "knowledge_bases.json"
    with open(kb_file, "w", encoding="utf-8") as f:
        json.dump(knowledge_bases, f, ensure_ascii=False, indent=2)


def read_file_content(file_path: str) -> str:
    """讀取不同類型文件的內容"""
    try:
        file_extension = Path(file_path).suffix.lower()

        if file_extension == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                return "\n".join(
                    page.extract_text() for page in pdf.pages if page.extract_text()
                )
        elif file_extension == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif file_extension == ".docx":
            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError(f"不支援的檔案類型: {file_extension}")
    except Exception as e:
        print(f"讀取檔案時出錯: {str(e)}")
        raise


def allowed_file(filename: str) -> bool:
    return (
        "." in filename
        and filename.rsplit(".", 1)[1].lower() in Config.ALLOWED_EXTENSIONS
    )


# 知識庫相關的API路由
@app.get("/api/knowledge_bases")
async def get_knowledge_bases():
    knowledge_bases = load_knowledge_bases()
    return [
        {"id": k, "name": v["name"], "description": v["description"]}
        for k, v in knowledge_bases.items()
    ]


@app.post("/api/knowledge_base")
async def create_knowledge_base(kb: KnowledgeBase):
    try:
        kb_id = str(uuid.uuid4())
        kb_path = Path(Config.CHROMA_PATH) / kb_id
        kb_path.mkdir(parents=True, exist_ok=True)

        knowledge_bases = load_knowledge_bases()
        knowledge_bases[kb_id] = {
            "name": kb.name,
            "description": kb.description,
            "path": str(kb_path),
        }
        save_knowledge_bases(knowledge_bases)

        # 初始化新知識庫的向量存儲
        initialize_vector_store(str(kb_path))

        return {"id": kb_id, "name": kb.name, "description": kb.description}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/knowledge_base/{kb_id}")
async def delete_knowledge_base(kb_id: str):
    try:
        if kb_id == "default":
            raise HTTPException(status_code=400, detail="無法刪除預設知識庫")

        knowledge_bases = load_knowledge_bases()
        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        kb_path = Path(knowledge_bases[kb_id]["path"])

        # 1. 首先關閉當前的向量存儲連接
        global vector_store
        if kb_id == Config.current_kb_id:
            try:
                vector_store._client.close()
                vector_store = None
            except:
                pass
            # 重新初始化默認知識庫
            vector_store = initialize_vector_store(
                str(Path(Config.CHROMA_PATH) / "default")
            )
            Config.current_kb_id = "default"

        # 2. 嘗試安刪除目錄
        try:
            safely_delete_directory(kb_path)
        except Exception as e:
            print(f"刪除知識庫目錄失敗: {str(e)}")
            # 如果刪除失敗，我們至少更新配置
            pass

        # 3. 新配置文件
        del knowledge_bases[kb_id]
        save_knowledge_bases(knowledge_bases)

        # 4. 清理可能的臨時文件
        import gc

        gc.collect()  # 強制垃圾回收

        return {"success": True}
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"刪除知識庫時出: {str(e)}")


@app.post("/api/knowledge_base/switch/{kb_id}")
async def switch_knowledge_base(kb_id: str):
    try:
        knowledge_bases = load_knowledge_bases()
        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        global vector_store
        # 先嘗試關閉現有連接
        try:
            if vector_store is not None:
                vector_store._client.close()
        except:
            pass

        vector_store = initialize_vector_store(knowledge_bases[kb_id]["path"])
        Config.current_kb_id = kb_id

        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/knowledge_base/reset/{kb_id}")
async def reset_knowledge_base(kb_id: str):
    try:
        knowledge_bases = load_knowledge_bases()
        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        if kb_id == Config.current_kb_id:
            reset_vector_store(vector_store)
        else:
            temp_store = initialize_vector_store(knowledge_bases[kb_id]["path"])
            reset_vector_store(temp_store)
            # 清理臨時存儲
            try:
                temp_store._client.close()
            except:
                pass

        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/translate")
async def translate(request: TranslateRequest):
    try:
        translated_text = one_chunk_translate_text(
            request.text,
            Config.MODEL_NAME,
            Config.SOURCE_LANG,
            Config.TARGET_LANG,
            Config.COUNTRY,
        )
        return {"translated_text": translated_text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload")
async def upload_file(
    files: List[UploadFile] = File(...), path: str = Form(default="/")
):
    """上傳檔案"""
    try:
        base_path = Path(Config.UPLOAD_FOLDER)
        current_path = base_path / path.lstrip("/")
        current_path.mkdir(parents=True, exist_ok=True)

        uploaded_files = []
        for file in files:
            if not file.filename:
                continue

            # 保存檔案
            file_path = current_path / file.filename
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            # 返回檔案資訊
            file_stat = file_path.stat()
            uploaded_files.append(
                {
                    "id": file.filename,  # 使用檔案名作為 ID
                    "name": file.filename,
                    "path": str(current_path.relative_to(base_path)),
                    "size": file_stat.st_size,
                    "type": file_path.suffix[1:] if file_path.suffix else "",
                    "created_at": datetime.fromtimestamp(
                        file_stat.st_ctime
                    ).isoformat(),
                    "isDirectory": False,
                }
            )

        return uploaded_files

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/upload_and_translate")
async def upload_and_translate(file: UploadFile = File(...)):
    """上傳並翻譯檔案"""
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="沒有提供文件")

    if not allowed_file(file.filename):
        raise HTTPException(status_code=400, detail="不允許的文件類型")

    try:
        filename = file.filename
        file_path = os.path.join(Config.UPLOAD_FOLDER, filename)

        # 保存上傳的文件
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        async def progress_callback(progress: int):
            await broadcast_progress(progress)

        translated_path = translate_and_store_to_knowledge_base(
            file_path,
            Config.MODEL_NAME,
            Config.SOURCE_LANG,
            Config.TARGET_LANG,
            Config.COUNTRY,
            progress_callback,
        )

        with open(translated_path, "r", encoding="utf-8") as f:
            translated_content = f.read()

        # 清理臨時文件
        os.remove(file_path)
        os.remove(translated_path)

        return {"translated_content": translated_content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"翻譯文件時出錯: {str(e)}")


@app.post("/api/embed")
async def embed_content(request: EmbedRequest):
    try:
        if not request.content:
            raise HTTPException(status_code=400, detail="沒有提供內容")

        kb_id = request.knowledge_base_id or Config.current_kb_id
        knowledge_bases = load_knowledge_bases()

        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        global vector_store
        if kb_id != Config.current_kb_id:
            temp_store = initialize_vector_store(knowledge_bases[kb_id]["path"])
            doc_id = add_translated_content_to_vector_store(
                temp_store,
                request.content,
                {
                    "source": request.filename,
                    "knowledge_base_id": kb_id,
                    "doc_id": str(uuid.uuid4()),
                },
            )
            # 清理臨時存儲
            try:
                temp_store._client.close()
            except:
                pass
        else:
            doc_id = add_translated_content_to_vector_store(
                vector_store,
                request.content,
                {
                    "source": request.filename,
                    "knowledge_base_id": kb_id,
                    "doc_id": str(uuid.uuid4()),
                },
            )

        return {"success": True, "doc_id": doc_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/knowledge_base/{kb_id}/file/{file_id}")
async def delete_file_from_knowledge_base(kb_id: str, file_id: str):
    try:
        knowledge_bases = load_knowledge_bases()
        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        global vector_store
        if kb_id != Config.current_kb_id:
            temp_store = initialize_vector_store(knowledge_bases[kb_id]["path"])
            success = delete_from_vector_store(temp_store, {"doc_id": file_id})
            # 清臨時存儲
            try:
                temp_store._client.close()
            except:
                pass
        else:
            success = delete_from_vector_store(vector_store, {"doc_id": file_id})

        if success:
            return {"success": True}
        else:
            raise HTTPException(status_code=404, detail="文件不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/query")
async def query(request: QueryRequest):
    try:
        kb_id = request.knowledge_base_id or Config.current_kb_id
        knowledge_bases = load_knowledge_bases()

        if kb_id not in knowledge_bases:
            raise HTTPException(status_code=404, detail="知識庫不存在")

        global vector_store
        current_vector_store = None

        if kb_id != Config.current_kb_id:
            current_vector_store = initialize_vector_store(
                knowledge_bases[kb_id]["path"]
            )
        else:
            current_vector_store = vector_store

        try:
            answer = query_knowledge_base(
                vector_store=current_vector_store,
                ffm=ffm,
                query=request.query,
                model_settings=request.model_settings,
            )

            # 獲取相關文件片段
            top_k = (
                request.model_settings.get("parameters", {}).get("topK", 3)
                if request.model_settings
                else 3
            )
            docs = current_vector_store.similarity_search(request.query, k=top_k)
            chunks = [doc.page_content for doc in docs]

            # 如果用了臨時向量存儲，清理它
            if kb_id != Config.current_kb_id and current_vector_store:
                try:
                    current_vector_store._client.close()
                except:
                    pass

            return {"answer": answer, "relevant_chunks": chunks}

        finally:
            # 確保在出現錯誤時能清理臨時向量存儲
            if kb_id != Config.current_kb_id and current_vector_store:
                try:
                    current_vector_store._client.close()
                except:
                    pass

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# 新增檔案管理相關的路由
@app.get("/api/files")
async def get_files(path: str = "/", search: Optional[str] = None):
    """獲取檔案和資料夾列表"""
    try:
        base_path = Path(Config.UPLOAD_FOLDER)
        current_path = base_path / path.lstrip("/")

        # 檢查路徑是否在 UPLOAD_FOLDER 內
        if not str(current_path.resolve()).startswith(str(base_path.resolve())):
            raise HTTPException(status_code=400, detail="無效的路徑")

        if not current_path.exists():
            return []

        items = []
        # 添加資料夾
        for item in current_path.iterdir():
            if search and search.lower() not in item.name.lower():
                continue

            stat = item.stat()
            relative_path = str(item.relative_to(base_path))

            if item.is_dir():
                items.append(
                    {
                        "id": relative_path,
                        "name": item.name,
                        "path": relative_path,
                        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        "isDirectory": True,
                        "size": 0,
                        "type": "folder",
                    }
                )
            else:
                items.append(
                    {
                        "id": relative_path,
                        "name": item.name,
                        "path": relative_path,
                        "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                        "isDirectory": False,
                        "size": stat.st_size,
                        "type": item.suffix[1:] if item.suffix else "",
                    }
                )

        return items
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/files/upload")
async def upload_file(
    files: List[UploadFile] = File(...), path: str = Form(default="/")
):
    """上傳檔案"""
    try:
        base_path = Path(Config.UPLOAD_FOLDER)
        current_path = base_path / path.lstrip("/")

        # 確保目標目錄存在
        current_path.mkdir(parents=True, exist_ok=True)

        uploaded_files = []
        for file in files:
            if not file.filename:
                continue

            file_path = current_path / file.filename

            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)

            file_stat = file_path.stat()
            uploaded_files.append(
                {
                    "id": str(file_path.relative_to(base_path)),
                    "name": file.filename,
                    "path": str(current_path.relative_to(base_path)),
                    "size": file_stat.st_size,
                    "type": file_path.suffix[1:] if file_path.suffix else "",
                    "created_at": datetime.fromtimestamp(
                        file_stat.st_ctime
                    ).isoformat(),
                    "isDirectory": False,
                }
            )

        return uploaded_files
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/files/{file_id}")
async def delete_file(file_id: str):
    """刪除檔案"""
    try:
        file_path = Path(Config.UPLOAD_FOLDER) / file_id
        if file_path.exists():
            file_path.unlink()
            return {"message": "檔案刪除成功"}
        raise HTTPException(status_code=404, detail="檔案不存在")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# 新增創建資夾的請求模型
class CreateFolderRequest(BaseModel):
    path: str


@app.post("/api/files/create_folder")
async def create_folder(request: CreateFolderRequest):
    """創建資料夾"""
    try:
        # 確保路徑是相對於 UPLOAD_FOLDER 的
        folder_path = Path(Config.UPLOAD_FOLDER) / request.path.lstrip("/")

        # 檢查路徑是否在 UPLOAD_FOLDER 內
        if not str(folder_path.resolve()).startswith(
            str(Path(Config.UPLOAD_FOLDER).resolve())
        ):
            raise HTTPException(status_code=400, detail="無效的路徑")

        # 創建資料夾
        folder_path.mkdir(parents=True, exist_ok=True)

        # 返回資料夾資訊
        folder_stat = folder_path.stat()
        return {
            "id": str(folder_path.relative_to(Config.UPLOAD_FOLDER)),
            "name": folder_path.name,
            "created_at": datetime.fromtimestamp(folder_stat.st_ctime).isoformat(),
            "isDirectory": True,
        }
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/api/files/folder/{folder_path:path}")
async def delete_folder(folder_path: str):
    """刪除資料夾"""
    try:
        folder_path = Path(Config.UPLOAD_FOLDER) / folder_path.lstrip("/")

        # 檢查路徑是否在 UPLOAD_FOLDER 內
        if not str(folder_path.resolve()).startswith(
            str(Path(Config.UPLOAD_FOLDER).resolve())
        ):
            raise HTTPException(status_code=400, detail="無效的路徑")

        if not folder_path.exists():
            raise HTTPException(status_code=404, detail="資料夾不存在")

        if not folder_path.is_dir():
            raise HTTPException(status_code=400, detail="指定路徑不是資料夾")

        # 刪除資料夾及其內容
        shutil.rmtree(folder_path)
        return {"message": "資料夾刪除成功"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/files/content/{file_path:path}")
async def get_file_content(file_path: str):
    """獲取檔案內容"""
    try:
        # 確保 UPLOAD_FOLDER 存在
        upload_folder = Path(Config.UPLOAD_FOLDER)
        upload_folder.mkdir(parents=True, exist_ok=True)

        # 處理檔案路徑
        file_path = upload_folder / file_path.lstrip("/").replace("\\", "/")
        print(f"嘗試讀取檔案: {file_path}")  # 添加日誌

        if not file_path.exists():
            print(f"檔案不存在: {file_path}")  # 添加日誌
            raise HTTPException(status_code=404, detail="檔案不存在")

        # 讀取檔案內容
        file_extension = file_path.suffix.lower()
        print(f"檔案類型: {file_extension}")  # 添加日誌

        try:
            if file_extension == ".pdf":
                with pdfplumber.open(str(file_path)) as pdf:
                    content = "\n".join(
                        page.extract_text() for page in pdf.pages if page.extract_text()
                    )
            elif file_extension == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            elif file_extension == ".docx":
                doc = Document(file_path)
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            else:
                raise HTTPException(
                    status_code=400, detail=f"不支援的檔案類型: {file_extension}"
                )

            return {"content": content, "filename": file_path.name}

        except Exception as e:
            print(f"讀取檔案時出錯: {str(e)}")  # 添加日誌
            raise HTTPException(status_code=500, detail=f"讀取檔案時出錯: {str(e)}")

    except HTTPException as e:
        raise e
    except Exception as e:
        print(f"處理請求時出錯: {str(e)}")  # 添加日誌
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/test-file-content")
async def test_file_content():
    """測試檔案內容讀取"""
    try:
        # 創建一個測試文件
        test_file_path = Path(Config.UPLOAD_FOLDER) / "test.txt"
        with open(test_file_path, "w", encoding="utf-8") as f:
            f.write("這是一個測試文件")

        # 讀取文件內容
        content = read_file_content(str(test_file_path))

        # 清理測試文件
        test_file_path.unlink()

        return {"content": content}
    except Exception as e:
        print(f"測試時出錯: {str(e)}")  # 添加詳細的錯誤日誌
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/test-file-read")
async def test_file_read():
    """測試檔案讀取功能"""
    try:
        # 創建一個測試檔案
        test_content = "This is a test file content."
        test_file_path = Path(Config.UPLOAD_FOLDER) / "test.txt"

        # 確保目錄存在
        test_file_path.parent.mkdir(parents=True, exist_ok=True)

        # 寫入測試檔案
        with open(test_file_path, "w", encoding="utf-8") as f:
            f.write(test_content)

        print(f"測試檔案已創建: {test_file_path}")

        # 嘗試讀取檔案
        try:
            content = read_file_content(str(test_file_path))
            print(f"成功讀取檔案內容: {content}")
        except Exception as e:
            print(f"讀取檔案時出錯: {str(e)}")
            raise

        # 清理測試檔案
        test_file_path.unlink()

        return {"success": True, "content": content}
    except Exception as e:
        print(f"測試時出錯: {str(e)}")
        return {"success": False, "error": str(e)}


if __name__ == "__main__":
    import uvicorn

    # 確保必要的目錄存在
    if not os.path.exists(Config.UPLOAD_FOLDER):
        os.makedirs(Config.UPLOAD_FOLDER)

    # 確保默認知識庫目錄存在
    default_kb_path = Path(Config.CHROMA_PATH) / "default"
    default_kb_path.mkdir(parents=True, exist_ok=True)

    # 初始化知識庫配置文件
    if not (Path(Config.CHROMA_PATH) / "knowledge_bases.json").exists():
        save_knowledge_bases(
            {
                "default": {
                    "name": "預設知識庫",
                    "description": "預設的知識庫",
                    "path": str(default_kb_path),
                }
            }
        )

    # 創建 translations 目錄（如果需要）
    translations_path = Path("translations")
    if not translations_path.exists():
        translations_path.mkdir(parents=True)

    uvicorn.run(app, host="0.0.0.0", port=5000)
